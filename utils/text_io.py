from __future__ import annotations
import io, json, csv, hashlib
from pathlib import Path
from typing import Tuple, Dict, List, Optional

from bs4 import BeautifulSoup
from charset_normalizer import from_bytes

try:
    from pypdf import PdfReader
except Exception:  # optional dep
    PdfReader = None

try:
    from docx import Document
except Exception:
    Document = None


def detect_encoding(data: bytes) -> str:
    """Guess encoding for unknown text files."""
    try:
        best = from_bytes(data).best()
        return best.encoding or "utf-8"
    except Exception:
        return "utf-8"


def sha256_12(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()[:12]


def chunk_text(text: str, max_chars: int = 12000) -> List[str]:
    """Simple char-based chunking (~3–4 chars per token)."""
    return [text[i:i+max_chars] for i in range(0, len(text), max_chars)]


def extract_text_from_bytes(name: str, data: bytes) -> Tuple[str, Dict]:
    """
    Convert uploaded file bytes to plain text.
    Supports: .txt .md .log .json .csv .html .htm .pdf .docx
    Raises ValueError with a friendly message on failure.
    """
    ext = Path(name).suffix.lower()
    meta: Dict = {
        "filename": name,
        "size_bytes": len(data),
        "type": ext.lstrip("."),
        "sha256_12": sha256_12(data),
        "pages": None,
        "encoding": None,
    }

    if ext in {".txt", ".md", ".log"}:
        enc = detect_encoding(data)
        meta["encoding"] = enc
        return data.decode(enc, errors="ignore"), meta

    if ext == ".json":
        enc = detect_encoding(data)
        meta["encoding"] = enc
        obj = json.loads(data.decode(enc, errors="ignore"))
        return json.dumps(obj, ensure_ascii=False, indent=2), meta

    if ext == ".csv":
        enc = detect_encoding(data)
        meta["encoding"] = enc
        s = data.decode(enc, errors="ignore")
        rows = list(csv.reader(io.StringIO(s)))
        lines = ["\t".join(row) for row in rows]
        return "\n".join(lines), meta

    if ext in {".html", ".htm"}:
        soup = BeautifulSoup(data, "lxml")
        for t in soup.find_all(["script", "style", "footer", "nav", "aside"]):
            t.decompose()
        body = soup.body or soup
        return body.get_text(separator="\n\n", strip=True), meta

    if ext == ".pdf":
        if PdfReader is None:
            raise ValueError("PDF parsing requires 'pypdf'. Install: pip install pypdf")
        reader = PdfReader(io.BytesIO(data))
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")  # sometimes empty password works
            except Exception:
                pass
        pages = [(p.extract_text() or "").strip() for p in reader.pages]
        text = "\n\n".join(p for p in pages if p)
        meta["pages"] = len(reader.pages)
        if not text.strip():
            raise ValueError("PDF has no extractable text (likely scanned). Use OCR.")
        return text, meta

    if ext == ".docx":
        if Document is None:
            raise ValueError("DOCX parsing requires 'python-docx'. Install: pip install python-docx")
        doc = Document(io.BytesIO(data))
        paras = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n\n".join(paras)
        if not text.strip() and doc.tables:
            rows = []
            for tbl in doc.tables:
                for row in tbl.rows:
                    rows.append("\t".join(c.text for c in row.cells))
            text = "\n".join(rows)
        if not text.strip():
            raise ValueError("DOCX contained no readable text. Try exporting to PDF.")
        return text, meta

    raise ValueError(f"Unsupported file type: {ext}")
